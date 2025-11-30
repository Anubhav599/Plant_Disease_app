from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Create a Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_paragraph()
title_run = title.add_run("PLANT DISEASE CLASSIFIER")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 71, 136)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("AI/ML Model - Complete Technical Documentation")
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(37, 99, 235)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Metadata
doc.add_paragraph()
meta = doc.add_paragraph()
meta_run = meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
meta_run.font.size = Pt(10)
meta_run.italic = True

# ================== SECTION 1: PROJECT OVERVIEW ==================
doc.add_heading("1. PROJECT OVERVIEW", level=1)

overview_text = """A Plant Disease Classification System using Deep Learning (Convolutional Neural Networks) to identify 38 different plant diseases from leaf images. The system uses TensorFlow/Keras for model training and Streamlit for web-based deployment. The model processes leaf images and provides real-time disease classification with probability scores for accurate plant health diagnosis."""
doc.add_paragraph(overview_text, style='List Bullet')

# Key Statistics Table
doc.add_heading("Key Project Statistics", level=2)
table = doc.add_table(rows=10, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Value'

# Data rows
data = [
    ['Total Classes', '38 plant diseases'],
    ['Dataset Source', 'PlantVillage (Kaggle)'],
    ['Training Samples', '54,305 images'],
    ['Input Image Size', '224×224 pixels (RGB)'],
    ['Model Architecture', 'Convolutional Neural Network (CNN)'],
    ['Training Epochs', '5'],
    ['Batch Size', '32'],
    ['Train/Validation Split', '80%/20%'],
    ['Model Format', 'HDF5 (.h5)'],
]

for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

doc.add_paragraph()

# ================== SECTION 2: MACHINE LEARNING CONCEPTS ==================
doc.add_heading("2. MACHINE LEARNING CONCEPTS", level=1)

doc.add_heading("2.1 Deep Learning Architecture - CNN (Convolutional Neural Network)", level=2)

cnn_intro = """A Convolutional Neural Network (CNN) is a specialized deep learning architecture designed for processing gridded data like images. It automatically learns spatial hierarchies of features through convolutional layers, making it ideal for image classification tasks."""
doc.add_paragraph(cnn_intro)

doc.add_heading("CNN Architecture Layers", level=3)

# CNN Layers Table
cnn_table = doc.add_table(rows=8, cols=3)
cnn_table.style = 'Light Grid Accent 1'

cnn_header = cnn_table.rows[0].cells
cnn_header[0].text = 'Layer'
cnn_header[1].text = 'Parameters'
cnn_header[2].text = 'Purpose'

cnn_data = [
    ['Conv2D (Layer 1)', '32 filters, 3×3 kernel', 'Extract low-level features (edges, textures)'],
    ['MaxPooling2D', '2×2 pool size', 'Reduce spatial dimensions, retain features'],
    ['Conv2D (Layer 2)', '64 filters, 3×3 kernel', 'Extract higher-level features (patterns, shapes)'],
    ['MaxPooling2D', '2×2 pool size', 'Further dimensionality reduction'],
    ['Flatten', 'N/A', 'Convert 2D feature maps to 1D vector'],
    ['Dense', '256 neurons, ReLU activation', 'Learn non-linear combinations of features'],
    ['Output Dense', '38 neurons, Softmax activation', '38-class probability distribution'],
]

for i, row_data in enumerate(cnn_data, 1):
    row_cells = cnn_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

doc.add_paragraph()

doc.add_heading("Why CNN?", level=3)
why_cnn_points = [
    "Automatically learn spatial hierarchies of features",
    "Exhibit shift and rotation invariance through convolution operations",
    "Share parameters across the image reducing model complexity",
    "Require fewer parameters than fully connected networks",
    "Perform exceptionally well on image recognition tasks",
]

for point in why_cnn_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

doc.add_heading("2.2 Key Machine Learning Concepts", level=2)

ml_table = doc.add_table(rows=11, cols=3)
ml_table.style = 'Light Grid Accent 1'

ml_header = ml_table.rows[0].cells
ml_header[0].text = 'Concept'
ml_header[1].text = 'Explanation'
ml_header[2].text = 'Application'

ml_data = [
    ['Convolution', 'Applies filters across image to detect patterns', 'Feature extraction from leaf images'],
    ['Pooling', 'Reduces dimensionality while preserving features', 'Prevents overfitting, improves efficiency'],
    ['Activation Function (ReLU)', 'Introduces non-linearity (max(0, x))', 'Enables learning complex boundaries'],
    ['Softmax', 'Converts logits to probability distribution', 'Multi-class classification output'],
    ['Categorical Cross-Entropy', 'Loss function for multi-class problems', 'Measures prediction error'],
    ['Optimization (Adam)', 'Adaptive learning rate optimizer', 'Efficiently updates network weights'],
    ['Epochs', 'Complete passes through training data', '5 epochs used in this model'],
    ['Batch Size', 'Number of samples processed together', '32 samples per batch'],
    ['Validation Split', 'Percentage of data for validation', '20% used for testing generalization'],
    ['Normalization', 'Scaling pixel values to [0, 1] range', 'Improves training stability and speed'],
]

for i, row_data in enumerate(ml_data, 1):
    row_cells = ml_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

doc.add_page_break()

# ================== SECTION 3: IMAGE PREPROCESSING ==================
doc.add_heading("3. IMAGE PREPROCESSING PIPELINE", level=1)

preprocess_intro = """Image preprocessing is a critical step that prepares raw input data for the neural network. The pipeline transforms raw images into a format optimized for the CNN model."""
doc.add_paragraph(preprocess_intro)

doc.add_heading("Preprocessing Sequence", level=2)
pipeline_steps = [
    "Input Image (JPG/PNG)",
    "Load using PIL (Python Imaging Library)",
    "Resize to 224×224 pixels",
    "Convert to NumPy array",
    "Add batch dimension: (1, 224, 224, 3)",
    "Normalize: divide by 255 to scale to [0, 1]",
    "Forward pass through CNN model",
]

for i, step in enumerate(pipeline_steps, 1):
    doc.add_paragraph(f"Step {i}: {step}", style='List Number')

doc.add_paragraph()

doc.add_heading("3.1 Preprocessing Steps (Detailed)", level=2)

preprocess_table = doc.add_table(rows=7, cols=3)
preprocess_table.style = 'Light Grid Accent 1'

preprocess_header = preprocess_table.rows[0].cells
preprocess_header[0].text = 'Step'
preprocess_header[1].text = 'Operation'
preprocess_header[2].text = 'Reason/Purpose'

preprocess_data = [
    ['1. Load Image', 'PIL.Image opens JPG/PNG file', 'Initial input from user or disk'],
    ['2. Resize', 'Scale to 224×224 pixels', 'Standard CNN input size; consistent dimensions'],
    ['3. Convert to Array', 'NumPy array format', 'Required for tensor operations in TensorFlow'],
    ['4. Add Batch Dimension', 'Shape: (224, 224, 3) → (1, 224, 224, 3)', 'Model expects batch of images'],
    ['5. Normalize', 'Divide pixel values by 255', 'Scale to [0, 1]; prevents gradient explosion'],
    ['6. Forward Pass', 'Feed through CNN layers', 'Extract features and classify'],
]

for i, row_data in enumerate(preprocess_data, 1):
    row_cells = preprocess_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

doc.add_paragraph()

why_normalize = doc.add_paragraph()
why_normalize_run = why_normalize.add_run("Why Normalization? ")
why_normalize_run.bold = True
normalize_text = "Networks train significantly better with normalized inputs because: (1) Prevents vanishing/exploding gradients, (2) Accelerates training convergence, (3) Improves numerical stability, (4) Reduces internal covariate shift, (5) Standardizes input scale across all features."
why_normalize.add_run(normalize_text)

doc.add_page_break()

# ================== SECTION 4: MODEL TRAINING FLOW ==================
doc.add_heading("4. MODEL TRAINING FLOW (DETAILED)", level=1)

training_intro = """The training process involves multiple stages from data preparation through model optimization and evaluation. Each stage is critical for creating an accurate, generalizable model."""
doc.add_paragraph(training_intro)

# Training Flow Stages
training_stages = [
    ('1. DATA CURATION', 
     ['Download PlantVillage dataset from Kaggle (54,305 images)',
      'Extract 38 disease classes',
      'Use "color" variant images for better feature extraction']),
    
    ('2. DATA PREPARATION',
     ['Create ImageDataGenerator for preprocessing',
      'Split: 80% training, 20% validation',
      'Batch size: 32 samples',
      'Target size: 224×224 pixels',
      'Rescale: divide by 255 to normalize']),
    
    ('3. MODEL ARCHITECTURE',
     ['Build Sequential model (linear layer stack)',
      'Conv2D(32) + MaxPooling → Conv2D(64) + MaxPooling',
      'Flatten → Dense(256, ReLU) → Output Dense(38, Softmax)',
      'Total parameters: ~1.3M']),
    
    ('4. COMPILATION',
     ['Optimizer: Adam (adaptive moment estimation)',
      'Loss Function: Categorical cross-entropy',
      'Metrics: Accuracy for monitoring performance']),
    
    ('5. TRAINING',
     ['Forward pass: Input → through all layers → output probabilities',
      'Backward propagation: Calculate gradients',
      'Weight updates: Adam optimizer adjusts all parameters',
      'Validation: Evaluate on held-out validation data',
      'Repeated for 5 epochs (5 complete passes through dataset)']),
    
    ('6. EVALUATION & STORAGE',
     ['Calculate validation accuracy',
      'Plot training vs validation accuracy curves',
      'Plot training vs validation loss curves',
      'Save model as plant_disease_model.h5 (HDF5 format)'])
]

for stage_title, stage_points in training_stages:
    doc.add_heading(stage_title, level=2)
    for point in stage_points:
        doc.add_paragraph(point, style='List Bullet')
    doc.add_paragraph()

doc.add_page_break()

# ================== SECTION 5: INFERENCE PIPELINE ==================
doc.add_heading("5. INFERENCE PIPELINE (DEPLOYMENT - main.py)", level=1)

inference_intro = """The inference pipeline is used when the model is deployed and makes predictions on new, unseen data. It follows the sequence: User Upload Image → Load & Preprocess → Model Prediction → Class Mapping → Display Result"""
doc.add_paragraph(inference_intro)

# Detailed Inference Steps
inference_steps = [
    ('1. STREAMLIT UI INITIALIZATION',
     ['Display title: "Plant Disease Classifier"',
      'Create file uploader widget',
      'Accept: JPG, JPEG, PNG formats']),
    
    ('2. MODEL & CLASS LOADING',
     ['Load pre-trained model: plant_disease_model.h5',
      'Load class indices: class_indices.json (38 disease mappings)',
      'Model loaded once during app startup for efficiency']),
    
    ('3. IMAGE UPLOAD',
     ['User selects leaf image through web interface',
      'Image stored in temporary memory',
      'No direct file system access required']),
    
    ('4. PREPROCESSING',
     ['Open image using PIL',
      'Resize to 224×224 pixels',
      'Convert to NumPy array',
      'Add batch dimension: (1, 224, 224, 3)',
      'Normalize: divide by 255 → [0, 1] range']),
    
    ('5. MODEL INFERENCE',
     ['Forward pass through 7 layers',
      'Output: 38 probability values (sum = 1.0)',
      'Computation time: ~100-500ms depending on hardware']),
    
    ('6. CLASS PREDICTION',
     ['Use np.argmax() to find highest probability index',
      'Map index to disease name using class_indices.json',
      'Example: Index 0 → "Apple___Apple_scab"']),
    
    ('7. RESULT DISPLAY',
     ['Display original image (150×150 px)',
      'Show prediction result using st.success()',
      'Format: "Prediction: [Disease_Name]"',
      'Probability scores could be added for confidence level'])
]

for step_title, step_points in inference_steps:
    doc.add_heading(step_title, level=2)
    for point in step_points:
        doc.add_paragraph(point, style='List Bullet')
    doc.add_paragraph()

doc.add_page_break()

# ================== SECTION 6: DATASET ==================
doc.add_heading("6. DATASET - 38 PLANT DISEASES", level=1)

dataset_intro = """The model is trained on the PlantVillage dataset containing 54,305 high-quality leaf images across 14 crop types and 38 disease/health classes."""
doc.add_paragraph(dataset_intro)

# Dataset Table
doc.add_heading("Complete Classification Scheme", level=2)
dataset_table = doc.add_table(rows=16, cols=3)
dataset_table.style = 'Light Grid Accent 1'

dataset_header = dataset_table.rows[0].cells
dataset_header[0].text = 'Crop'
dataset_header[1].text = 'Classes'
dataset_header[2].text = 'Count'

dataset_data = [
    ['Apple', 'Apple_scab, Black_rot, Cedar_apple_rust, healthy', '4'],
    ['Blueberry', 'healthy', '1'],
    ['Cherry', 'Powdery_mildew, healthy', '2'],
    ['Corn (Maize)', 'Cercospora, Common_rust, Northern_Leaf_Blight, healthy', '4'],
    ['Grape', 'Black_rot, Esca, Leaf_blight, healthy', '4'],
    ['Orange', 'Haunglongbing (Citrus_greening)', '1'],
    ['Peach', 'Bacterial_spot, healthy', '2'],
    ['Pepper, Bell', 'Bacterial_spot, healthy', '2'],
    ['Potato', 'Early_blight, Late_blight, healthy', '3'],
    ['Raspberry', 'healthy', '1'],
    ['Soybean', 'healthy', '1'],
    ['Squash', 'Powdery_mildew', '1'],
    ['Strawberry', 'Leaf_scorch, healthy', '2'],
    ['Tomato', 'Bacterial_spot, Early_blight, Late_blight, Leaf_Mold, Septoria_leaf_spot, Spider_mites, Target_Spot, YLCV, TMV, healthy', '10'],
    ['TOTAL', '', '38'],
]

for i, row_data in enumerate(dataset_data, 1):
    row_cells = dataset_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

doc.add_page_break()

# ================== SECTION 7: TECHNICAL STACK ==================
doc.add_heading("7. TECHNICAL STACK & DEPENDENCIES", level=1)

tech_table = doc.add_table(rows=9, cols=4)
tech_table.style = 'Light Grid Accent 1'

tech_header = tech_table.rows[0].cells
tech_header[0].text = 'Component'
tech_header[1].text = 'Technology'
tech_header[2].text = 'Version'
tech_header[3].text = 'Purpose'

tech_data = [
    ['Deep Learning Framework', 'TensorFlow/Keras', 'Latest', 'Model training & inference, layer abstractions'],
    ['Image Processing', 'PIL (Pillow)', 'Latest', 'Load, resize, manipulate images'],
    ['Numerical Computing', 'NumPy', 'Latest', 'Array operations, matrix computations'],
    ['Web UI Framework', 'Streamlit', 'Latest', 'Interactive web interface, rapid prototyping'],
    ['Data Source', 'PlantVillage Dataset', 'Kaggle', '54,305 plant leaf images, 38 classes'],
    ['Containerization', 'Docker', '20.10+', 'Consistent deployment across environments'],
    ['Configuration', 'TOML', 'Python standard', 'Server settings, browser configuration'],
    ['Serialization', 'HDF5 (.h5)', 'TensorFlow standard', 'Model architecture & weights storage'],
]

for i, row_data in enumerate(tech_data, 1):
    row_cells = tech_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
    row_cells[3].text = row_data[3]

doc.add_page_break()

# ================== SECTION 8: MODEL ARCHITECTURE ==================
doc.add_heading("8. MODEL ARCHITECTURE VISUALIZATION", level=1)

arch_text = doc.add_paragraph()
arch_flow = """
INPUT LAYER (224×224×3 RGB Image)
↓
Convolutional Layer 1 (Conv2D with 32 filters, 3×3 kernel) + ReLU Activation
↓
Max Pooling Layer 1 (2×2 pool size) [Reduces spatial dimensions to ~112×112]
↓
Convolutional Layer 2 (Conv2D with 64 filters, 3×3 kernel) + ReLU Activation
↓
Max Pooling Layer 2 (2×2 pool size) [Further reduces to ~56×56]
↓
Flatten Layer [Converts all feature maps to 1D vector]
↓
Fully Connected (Dense) Layer 1 (256 neurons, ReLU activation) [Non-linear feature combinations]
↓
Output Dense Layer (38 neurons, Softmax activation) [Probability distribution across diseases]
↓
OUTPUT [38 disease probability scores, sum = 1.0]
"""
arch_text.add_run(arch_flow).font.name = 'Courier New'

doc.add_paragraph()

doc.add_heading("Model Parameters Summary", level=2)
model_summary = doc.add_paragraph()
model_summary.add_run("Total Trainable Parameters: ").bold = True
model_summary.add_run("~1.3 Million\n")
model_summary.add_run("Convolutional Parameters: ").bold = True
model_summary.add_run("~900K\n")
model_summary.add_run("Dense Parameters: ").bold = True
model_summary.add_run("~400K")

doc.add_page_break()

# ================== SECTION 9: HYPERPARAMETERS ==================
doc.add_heading("9. HYPERPARAMETERS & CONFIGURATION", level=1)

hp_table = doc.add_table(rows=17, cols=3)
hp_table.style = 'Light Grid Accent 1'

hp_header = hp_table.rows[0].cells
hp_header[0].text = 'Parameter'
hp_header[1].text = 'Value'
hp_header[2].text = 'Rationale'

hp_data = [
    ['Image Input Size', '224×224 pixels', 'Standard for modern CNNs; balance between detail and computation'],
    ['Batch Size', '32 samples', 'Trade-off between memory usage and training stability'],
    ['Training Epochs', '5', 'Limited epochs in demo; production uses 50-200 for better convergence'],
    ['Validation Split', '20%', 'Sufficient for monitoring generalization without reducing training data'],
    ['Conv2D Filters (Layer 1)', '32', 'Captures low-level features; reduced complexity'],
    ['Conv2D Filters (Layer 2)', '64', 'Captures higher-level patterns; doubles filter count'],
    ['Kernel Size', '3×3', 'Standard for feature extraction; balance between receptive field and parameters'],
    ['Pooling Size', '2×2', 'Reduces spatial dimensions by 50% per pooling layer'],
    ['Dense Layer Neurons', '256', 'Sufficient capacity for learning feature combinations'],
    ['Output Layer Neurons', '38', 'One neuron per disease class'],
    ['ReLU Activation', 'max(0, x)', 'Introduces non-linearity; standard for hidden layers'],
    ['Softmax Activation', 'e^x / Σ(e^x)', 'Normalizes outputs to probabilities for multi-class'],
    ['Optimizer', 'Adam', 'Adaptive learning rate; faster convergence than SGD'],
    ['Learning Rate (Adam)', '~0.001 (default)', 'Balanced convergence speed and stability'],
    ['Loss Function', 'Categorical Cross-Entropy', 'Standard for multi-class classification'],
    ['Dropout (Optional)', 'Not used', 'Could add 0.2-0.5 dropout to prevent overfitting'],
]

for i, row_data in enumerate(hp_data, 1):
    row_cells = hp_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

doc.add_page_break()

# ================== SECTION 10: LOSS FUNCTIONS & OPTIMIZATION ==================
doc.add_heading("10. LOSS FUNCTIONS & OPTIMIZATION", level=1)

doc.add_heading("10.1 Categorical Cross-Entropy Loss", level=2)

loss_intro = """The loss function quantifies the difference between predicted and actual distributions. Categorical Cross-Entropy is the standard for multi-class classification."""
doc.add_paragraph(loss_intro)

doc.add_heading("Formula", level=3)
formula_p = doc.add_paragraph()
formula_run = formula_p.add_run("L = -Σ(y_i × log(ŷ_i))")
formula_run.bold = True
formula_run.font.size = Pt(12)

doc.add_heading("Where:", level=3)
where_points = [
    "y_i = true label (one-hot encoded, 0 or 1)",
    "ŷ_i = predicted probability for class i",
    "Σ = sum across all 38 classes",
    "log = natural logarithm",
]
for point in where_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

why_ce = doc.add_paragraph()
why_ce.add_run("Why Cross-Entropy? ").bold = True
why_ce.add_run("It heavily penalizes confident wrong predictions, encouraging the model to learn discriminative features. It has nice mathematical properties for backpropagation and gradient descent optimization.")

doc.add_paragraph()

doc.add_heading("10.2 Adam Optimizer", level=2)

adam_intro = """Adam (Adaptive Moment Estimation) is an advanced optimization algorithm that adapts learning rates for each parameter."""
doc.add_paragraph(adam_intro)

doc.add_heading("Key Features:", level=3)
adam_features = [
    "Maintains exponential moving average of gradients (momentum term)",
    "Maintains exponential moving average of squared gradients (RMSprop term)",
    "Adapts learning rate per parameter based on these moving averages",
    "Default learning rate: 0.001 (typically requires no tuning)",
    "Computationally efficient and memory-friendly",
]
for feature in adam_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()

doc.add_heading("Advantages over Standard SGD:", level=3)
sgd_advantages = [
    "Faster convergence with fewer iterations",
    "Handles sparse gradients well",
    "Robust to different hyperparameter choices",
    "Works well with mini-batches (batch size: 32)",
    "Reduces need for learning rate scheduling",
]
for adv in sgd_advantages:
    doc.add_paragraph(adv, style='List Bullet')

doc.add_page_break()

# ================== SECTION 11: DEPLOYMENT ARCHITECTURE ==================
doc.add_heading("11. DEPLOYMENT ARCHITECTURE (DOCKER)", level=1)

deploy_intro = """The application is containerized using Docker for consistent deployment across different environments. The container includes all dependencies, the trained model, and configuration files."""
doc.add_paragraph(deploy_intro)

doc.add_heading("Docker Container Structure:", level=2)

docker_points = [
    "Base Image: python:3.10-slim (lightweight Python runtime)",
    "COPY: Copy all application files to /app directory",
    "WORKDIR: Set /app as working directory",
    "RUN: Execute pip install for dependencies",
    "EXPOSE: Port 80 (HTTP traffic)",
    "Create ~/.streamlit directory",
    "Copy config.toml (server settings)",
    "Copy credentials.toml (authentication)",
    "ENTRYPOINT: streamlit run",
    "CMD: main.py",
]
for point in docker_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

doc.add_heading("Runtime Configuration:", level=2)

runtime_points = [
    "Starts Streamlit server on 0.0.0.0:80",
    "Accessible from any network interface",
    "Port 80: Standard HTTP port (no port forwarding needed)",
    "Server address: 0.0.0.0 (listens on all interfaces)",
    "CORS disabled for security",
    "Headless mode enabled",
    "Live save disabled for stability",
    "Run on save disabled (prevents auto-refresh)",
]
for point in runtime_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

doc.add_heading("Advantages:", level=2)

docker_advantages = [
    "Reproducible deployments across servers",
    "Isolation from host system dependencies",
    "Easy scaling and orchestration",
    "Version control for entire environment",
    "One-command deployment: docker run",
    "Works on Linux, Windows, macOS",
]
for adv in docker_advantages:
    doc.add_paragraph(adv, style='List Bullet')

doc.add_page_break()

# ================== SECTION 12: CRITICAL ML CONCEPTS ==================
doc.add_heading("12. CRITICAL ML CONCEPTS FOR DEVELOPMENT", level=1)

critical_concepts = [
    ('Overfitting Prevention',
     'Problem: Model memorizes training data instead of learning generalizable patterns.\n\nSolution: Use validation set (20%) to monitor generalization. If val_loss increases while train_loss decreases = overfitting.\n\nTechniques: Early stopping, dropout layers, regularization, data augmentation.'),
    
    ('Data Augmentation',
     'Creates variations of training images (rotations, flips, zooms, color changes).\n\nBenefits: Increases effective training data size, improves model robustness, prevents overfitting.\n\nImplementation: ImageDataGenerator in current code supports multiple augmentation techniques.'),
    
    ('Feature Extraction',
     'CNN layers learn hierarchical features: Layer 1 (edges) → Layer 2 (textures) → Dense layers (object parts).\n\nTransfer Learning: Use pre-trained weights (ImageNet) instead of training from scratch for faster convergence.'),
    
    ('Vanishing Gradients',
     'Problem: Gradients become too small during backpropagation in deep networks.\n\nSolution: ReLU activation, batch normalization, proper weight initialization.\n\nImpact: Without mitigation, network cannot learn effectively.'),
    
    ('Model Serialization',
     'Save/load models for reuse without retraining. HDF5 format (.h5) stores:\n• Model architecture (layer definitions)\n• Weights (learned parameters)\n• Training configuration\n\nBenefit: Inference uses pre-trained weights, ~1000x faster than training.'),
    
    ('Gradient Descent & Backpropagation',
     'Forward Pass: Input → layers → output predictions\n\nBackward Pass: Calculate ∂L/∂w (gradient of loss w.r.t. weights)\n\nUpdate: w_new = w_old - α × ∂L/∂w (where α = learning rate)\n\nRepeated for all layers in reverse order (hence "back" propagation).'),
]

for concept_title, concept_details in critical_concepts:
    doc.add_heading(concept_title, level=2)
    doc.add_paragraph(concept_details)
    doc.add_paragraph()

doc.add_page_break()

# ================== SECTION 13: PERFORMANCE METRICS ==================
doc.add_heading("13. PERFORMANCE METRICS & EVALUATION", level=1)

metrics_intro = """The model's performance is evaluated using multiple metrics to ensure accuracy and generalization."""
doc.add_paragraph(metrics_intro)

doc.add_paragraph()

metrics = [
    ('Accuracy',
     'Percentage of correct predictions: (True Positives + True Negatives) / Total\n\nRange: 0-100% (higher is better)\n\nUsage: Overall performance metric across all classes.'),
    
    ('Loss',
     'Categorical Cross-Entropy loss value\n\nRange: 0 to ∞ (lower is better)\n\nTraining Loss: Monitors learning on training data\n\nValidation Loss: Monitors generalization on unseen data\n\nIf val_loss > train_loss consistently → overfitting.'),
    
    ('Training vs Validation Curves',
     'Plotted against epochs to visualize learning\n\nIdeal: Both curves decrease smoothly and converge\n\nWarning Signs: Validation plateau while training continues = overfitting\n\nMetric in Code: Used to detect model performance trends.'),
    
    ('Confusion Matrix (Optional)',
     'Shows true/false positives/negatives per class\n\nIdentifies which diseases are confused with others\n\nUseful for understanding model weaknesses.\n\nCould enhance current implementation.'),
    
    ('Per-Class Metrics (Optional)',
     'Precision: TP / (TP + FP) - accuracy for each disease\n\nRecall: TP / (TP + FN) - disease detection rate\n\nF1-Score: Harmonic mean of precision & recall\n\nUseful for imbalanced datasets.'),
]

for metric_title, metric_details in metrics:
    doc.add_heading(metric_title, level=2)
    doc.add_paragraph(metric_details)

doc.add_page_break()

# ================== SECTION 14: COMPLETE WORKFLOW ==================
doc.add_heading("14. COMPLETE WORKFLOW SEQUENCE", level=1)

doc.add_heading("Training Phase (app.py):", level=2)

training_workflow = [
    'Download PlantVillage Dataset from Kaggle (54,305 images)',
    'Extract and organize 38 disease classes',
    'Create ImageDataGenerator with rescaling (÷255)',
    'Split data: 80% training, 20% validation',
    'Build Sequential CNN model (7 layers total)',
    'Compile with Adam optimizer & cross-entropy loss',
    'Train for 5 epochs with batch size 32',
    'Monitor training/validation accuracy and loss',
    'Evaluate on validation set',
    'Plot performance curves (accuracy & loss)',
    'Save model as plant_disease_model.h5 (HDF5 format)',
    'Save class indices mapping as class_indices.json',
]

for i, step in enumerate(training_workflow, 1):
    doc.add_paragraph(f'Step {i}: {step}', style='List Number')

doc.add_paragraph()

doc.add_heading("Inference Phase (main.py):", level=2)

inference_workflow = [
    'Load pre-trained model from plant_disease_model.h5',
    'Load class indices mapping from class_indices.json',
    'Initialize Streamlit web interface',
    'Display title and file uploader widget',
    'User uploads leaf image (JPG/JPEG/PNG)',
    'Display uploaded image (150×150 px preview)',
    'User clicks "Classify" button',
    'Preprocess image: resize to 224×224, normalize to [0,1]',
    'Run forward pass through CNN layers (~100-500ms)',
    'Model outputs 38 probability values',
    'Extract class with highest probability using argmax()',
    'Map class index to disease name',
    'Display prediction result to user',
    'Ready for next image classification',
]

for i, step in enumerate(inference_workflow, 1):
    doc.add_paragraph(f'Step {i}: {step}', style='List Number')

doc.add_page_break()

# ================== SECTION 15: FUTURE ENHANCEMENTS ==================
doc.add_heading("15. FUTURE ENHANCEMENTS & IMPROVEMENTS", level=1)

enhancements = [
    ('Transfer Learning',
     'Use pre-trained ImageNet weights as starting point instead of random initialization.\n\nBenefit: Faster training, better accuracy with fewer data, leverages general image features.\n\nExample: Use MobileNetV2 or ResNet50 as backbone.'),
    
    ('Data Augmentation',
     'Apply random rotations, flips, zoom, brightness changes to training images.\n\nBenefit: Simulates different lighting/angles, improves robustness.\n\nCurrent Status: Infrastructure exists (ImageDataGenerator) but not fully utilized.'),
    
    ('Confidence Scores',
     'Display probability for predicted class and top-3 alternatives.\n\nBenefit: Users understand model confidence, easier to validate predictions.\n\nImplementation: Extract top_3 indices from prediction array.'),
    
    ('Model Validation Metrics',
     'Add confusion matrix, precision/recall per class, ROC curves.\n\nBenefit: Detailed performance analysis, identify problematic classes.\n\nTools: scikit-learn confusion_matrix, classification_report.'),
    
    ('Batch Prediction',
     'Allow uploading multiple images at once.\n\nBenefit: Process folders of farm images, generate reports.\n\nImplementation: Loop through multiple files, collect predictions.'),
    
    ('Mobile Deployment',
     'Convert to TensorFlow Lite for mobile apps (Android/iOS).\n\nBenefit: On-device inference, no internet required, faster response.\n\nTools: tf.lite.TFLiteConverter.'),
    
    ('API Development',
     'Create REST API (FastAPI/Flask) for programmatic access.\n\nBenefit: Integration with other systems, batch processing workflows.\n\nEndpoint: POST /predict with image → returns JSON with classification.'),
    
    ('Explainability',
     'Visualize which image regions most influence predictions (Grad-CAM, LIME).\n\nBenefit: Build trust, understand model reasoning, debugging.\n\nTools: tf-explain, grad-cam libraries.'),
]

for enh_title, enh_details in enhancements:
    doc.add_heading(enh_title, level=2)
    doc.add_paragraph(enh_details)

doc.add_page_break()

# ================== SECTION 16: TROUBLESHOOTING ==================
doc.add_heading("16. COMMON ISSUES & TROUBLESHOOTING", level=1)

troubleshooting = [
    ('OutOfMemory Error',
     'Symptom: "CUDA out of memory" or "Cannot allocate memory"\n\nCauses: Batch size too large, model too large for GPU/RAM\n\nSolutions: Reduce batch size (try 16 or 8), use smaller model, enable gradient checkpointing.'),
    
    ('Model Accuracy Low (<70%)',
     'Symptom: Validation accuracy plateau at low value\n\nCauses: Insufficient training epochs, learning rate too high/low, bad data quality\n\nSolutions: Train longer (50+ epochs), adjust learning rate, verify image quality.'),
    
    ('Severe Overfitting',
     'Symptom: Train accuracy 95%+ but validation accuracy 50%\n\nCauses: Model too complex, insufficient data, no augmentation\n\nSolutions: Add dropout layers, use data augmentation, collect more data, reduce model size.'),
    
    ('Slow Inference',
     'Symptom: Predictions take 5+ seconds\n\nCauses: CPU inference (no GPU), model too large, disk I/O delays\n\nSolutions: Use GPU (CUDA), quantize model, optimize preprocessing.'),
    
    ('Image Preprocessing Errors',
     'Symptom: "Image size mismatch" or dimension errors\n\nCauses: Inconsistent image sizes, wrong color channels\n\nSolutions: Always resize to 224×224, ensure RGB format (not RGBA or grayscale).'),
    
    ('Class Index Mismatch',
     'Symptom: Predictions show wrong disease names\n\nCauses: class_indices.json doesn\'t match model training\n\nSolutions: Regenerate class_indices.json from train_generator.class_indices.'),
    
    ('Streamlit Not Loading',
     'Symptom: Page keeps loading, no error message\n\nCauses: Model loading takes too long, file path incorrect\n\nSolutions: Add caching (@st.cache), verify file paths, check console logs.'),
]

for issue_title, issue_details in troubleshooting:
    doc.add_heading(issue_title, level=2)
    doc.add_paragraph(issue_details)

doc.add_page_break()

# ================== SECTION 17: CONCLUSION ==================
doc.add_heading("17. CONCLUSION & KEY TAKEAWAYS", level=1)

doc.add_heading("Summary", level=2)

summary_text = """This Plant Disease Classifier demonstrates a complete end-to-end deep learning pipeline for practical agricultural applications. The system combines CNNs' powerful feature extraction capabilities with Streamlit's user-friendly interface for accessible plant disease diagnosis."""
doc.add_paragraph(summary_text)

doc.add_heading("Key Technical Achievements:", level=2)

achievements = [
    "Implemented 7-layer CNN architecture optimized for 38-class disease classification",
    "Preprocesses images to consistent format ([0,1] normalized, 224×224 pixels)",
    "Uses Adam optimizer with categorical cross-entropy loss for stable training",
    "Achieves reasonable accuracy within 5 training epochs",
    "Deployed via Docker for reproducible, scalable deployment",
    "Provides web interface for non-technical users",
    "Leverages Kaggle's PlantVillage dataset with 54,305 high-quality images",
]

for achievement in achievements:
    doc.add_paragraph(achievement, style='List Bullet')

doc.add_paragraph()

doc.add_heading("ML Concepts Demonstrated:", level=2)

concepts_list = [
    "Convolutional Neural Networks and feature hierarchies",
    "Data preprocessing and normalization techniques",
    "Train/validation split and overfitting prevention",
    "Backpropagation and gradient descent optimization",
    "Model serialization and inference pipelines",
    "Web interface development for ML models",
    "Containerization for consistent deployment",
]

for concept in concepts_list:
    doc.add_paragraph(concept, style='List Bullet')

doc.add_paragraph()

doc.add_heading("Real-World Applications:", level=2)

applications = [
    "Early disease detection for crop management",
    "Reduce pesticide waste through targeted treatment",
    "Prevent large-scale crop failures",
    "Enable data-driven agricultural decisions",
    "Lower farming costs and increase yields",
]

for app in applications:
    doc.add_paragraph(app, style='List Bullet')

doc.add_paragraph()

doc.add_heading("Next Steps for Production Deployment:", level=2)

next_steps = [
    'Collect more diverse training data from different farms/regions',
    'Implement comprehensive data augmentation',
    'Use transfer learning to improve accuracy',
    'Add confidence scores for predictions',
    'Create mobile app for field use',
    'Integrate with farm management systems',
    'Implement user feedback loop for continuous improvement',
    'Deploy on edge devices for offline operation',
]

for i, step in enumerate(next_steps, 1):
    doc.add_paragraph(f'Step {i}: {step}', style='List Number')

doc.add_paragraph()

doc.add_heading("Final Note", level=2)

final_note = """Understanding these technical details is crucial for maintaining, improving, and extending this system for practical agricultural applications. The modular architecture allows for easy integration of new components, testing of different architectures, and deployment in various environments."""
doc.add_paragraph(final_note)

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}").italic = True
footer.add_run("\n")
footer.add_run("Plant Disease Classifier - Technical Documentation").italic = True

# Save Document
doc.save('Plant_Disease_Classifier_Technical_Documentation.docx')
print("✓ Word Document created successfully!")
print("✓ File: Plant_Disease_Classifier_Technical_Documentation.docx")
print(f"✓ Location: {__file__.replace('generate_docx.py', '')}")
